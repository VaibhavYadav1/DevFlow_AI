import os
import markdown
import requests
import base64
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Inches

def generate_project_documents(markdown_content: str, mermaid_data: dict, output_dir: str):
    """
    Generates PDF and DOCX files from markdown content and mermaid diagrams.
    
    Args:
        markdown_content: The project summary in markdown format.
        mermaid_data: Dictionary containing mermaid diagram code (class_diagram, api_flow).
        output_dir: The directory to save the generated files.
    """
    
    os.makedirs(output_dir, exist_ok=True)
    
    # --- 1. Process Mermaid Diagrams ---
    diagram_images = {}
    
    for key, mermaid_code in mermaid_data.items():
        if not mermaid_code or mermaid_code.strip() == "":
            continue
            
        try:
            # Encode mermaid graph to base64 for mermaid.ink
            graph_bytes = mermaid_code.encode("utf8")
            base64_bytes = base64.b64encode(graph_bytes)
            base64_string = base64_bytes.decode("ascii")
            
            image_url = f"https://mermaid.ink/img/{base64_string}"
            
            # Download image for embedding
            response = requests.get(image_url)
            if response.status_code == 200:
                image_path = os.path.join(output_dir, f"{key}.png")
                with open(image_path, 'wb') as f:
                    f.write(response.content)
                diagram_images[key] = image_path
            else:
                print(f"Failed to download mermaid diagram from {image_url}")
                
        except Exception as e:
            print(f"Error processing mermaid diagram {key}: {e}")

    # --- 2. Generate PDF ---
    pdf_path = os.path.join(output_dir, "project_summary.pdf")
    try:
        html_content = markdown.markdown(markdown_content)
        
        # Append diagrams to HTML
        if diagram_images:
            html_content += "<h2>Architecture Diagrams</h2>"
            for key, image_path in diagram_images.items():
                 # Use absolute path for local file access in xhtml2pdf usually works better with absolute paths
                 abs_image_path = os.path.abspath(image_path)
                 html_content += f"<h3>{key.replace('_', ' ').title()}</h3>"
                 html_content += f'<img src="{abs_image_path}" style="max-width: 100%;">'

        # Add some basic styling
        full_html = f"""
        <html>
        <head>
            <style>
                body {{ font-family: Helvetica, sans-serif; font-size: 12px; }}
                h1 {{ color: #333; font-size: 24px; border-bottom: 2px solid #333; padding-bottom: 10px; margin-bottom: 20px; }}
                h2 {{ color: #555; font-size: 18px; margin-top: 20px; border-bottom: 1px solid #ddd; }}
                h3 {{ color: #777; font-size: 14px; margin-top: 15px; }}
                table {{ width: 100%; border-collapse: collapse; margin-top: 10px; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background-color: #f2f2f2; }}
                code {{ background-color: #f9f9f9; padding: 2px 4px; border-radius: 4px; font-family: monospace; }}
                pre {{ background-color: #f9f9f9; padding: 10px; border-radius: 4px; overflow-x: auto; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """

        with open(pdf_path, "wb") as pdf_file:
            pisa_status = pisa.CreatePDF(full_html, dest=pdf_file)
            
        if pisa_status.err:
            print(f"Error generating PDF: {pisa_status.err}")
            
    except Exception as e:
        print(f"Failed to generate PDF: {e}")

    # --- 3. Generate DOCX ---
    docx_path = os.path.join(output_dir, "project_summary.docx")
    try:
        doc = Document()
        doc.add_heading('Project Summary', 0)
        
        # Split markdown by double newlines to roughly approximate paragraphs
        lines = markdown_content.split('\n')
        for line in lines:
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.strip().startswith('- '):
                doc.add_paragraph(line.strip()[2:], style='List Bullet')
            elif line.strip():
                doc.add_paragraph(line)
                
        # Append Diagrams
        if diagram_images:
            doc.add_heading('Architecture Diagrams', level=1)
            for key, image_path in diagram_images.items():
                doc.add_heading(key.replace('_', ' ').title(), level=2)
                try:
                    doc.add_picture(image_path, width=Inches(6))
                except Exception as img_err:
                     doc.add_paragraph(f"[Image: {key} - could not be added: {img_err}]")

        doc.save(docx_path)
        
    except Exception as e:
        print(f"Failed to generate DOCX: {e}")
